from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

from .schemas import WordSchema
from .export_utils import build_export_path

def generate_word(schema: WordSchema) -> str:
    """
    Generate Word document from schema.
    Returns absolute file path.
    """

    doc = Document()

    # ----------------------------------
    # Header & Footer
    # ----------------------------------
    section = doc.sections[0]

    if schema.header_text:
        header = section.header
        header.paragraphs[0].text = schema.header_text

    if schema.footer_text:
        footer = section.footer
        footer.paragraphs[0].text = schema.footer_text

    # ----------------------------------
    # Paragraphs
    # ----------------------------------
    for para in schema.paragraphs:

        if para.page_break_before:
            doc.add_page_break()

        if para.heading_level is not None:
            level = max(0, min(para.heading_level, 9))
            p = doc.add_heading(para.text, level=level)

        elif para.list_type == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(para.text)

        elif para.list_type == "numbered":
            p = doc.add_paragraph(style="List Number")
            p.add_run(para.text)

        else:
            if para.style:
                p = doc.add_paragraph(style=para.style)
            else:
                p = doc.add_paragraph()

            run = p.add_run(para.text)
            run.bold = para.bold
            run.italic = para.italic

        # Alignment
        if para.alignment:
            alignment_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }
            p.alignment = alignment_map[para.alignment]

    # ----------------------------------
    # Tables
    # ----------------------------------
    for table_data in schema.tables:

        if not table_data.headers:
            continue

        table = doc.add_table(
            rows=1,
            cols=len(table_data.headers)
        )

        table.style = table_data.style

        # Header row
        header_cells = table.rows[0].cells
        for i, header in enumerate(table_data.headers):
            header_cells[i].text = header

            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Data rows
        for row in table_data.rows:
            row_cells = table.add_row().cells
            for i, cell in enumerate(row):
                row_cells[i].text = str(cell)

    # ----------------------------------
    # Images
    # ----------------------------------
    for image in schema.images:
        image_path = Path(image.path)

        if not image_path.exists():
            raise FileNotFoundError(
                f"Image not found: {image.path}"
            )

        if image.width:
            doc.add_picture(
                str(image_path),
                width=Inches(image.width)
            )
        else:
            doc.add_picture(str(image_path))

    # ----------------------------------
    # Export
    # ----------------------------------
    file_path = build_export_path(schema.file_name, "docx")
    doc.save(file_path)

    return str(file_path)
